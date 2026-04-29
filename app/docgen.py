"""
Stoken Advisory — Word Document Generator & Parser
Gera formulários .docx formatados e parseia formulários preenchidos.
"""

import io
import re
import logging
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

# ─── Perguntas por área (mesmas do frontend) ──────────────────────────────────

AREA_QUESTIONS = {
    "supply-chain": {
        "title": "Supply Chain",
        "questions": [
            "Como funciona o processo de S&OP atualmente? Existe validação estratégica de demanda?",
            "Qual o nível de integração entre as áreas (Comercial, PCP, Financeiro)?",
            "Como são definidos, acompanhados e revisados os KPIs da área?",
            "Quais são os maiores gargalos da cadeia de suprimentos hoje?",
            "Existe visibilidade end-to-end da cadeia? Quais ferramentas são usadas?",
            "Como é feita a gestão de riscos da cadeia (fornecedores, logística, demanda)?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "producao": {
        "title": "Produção / PCP",
        "questions": [
            "O planejamento de produção é feito em sistema ou em Excel? Por quê?",
            "Existe controle de acurácia do cronograma de produção?",
            "Como funciona o custeio de produtos? Qual a granularidade disponível?",
            "Qual o OEE médio das linhas? Como é medido?",
            "Como são tratadas as paradas não planejadas? Existe manutenção preventiva?",
            "Como é feito o sequenciamento de produção? Quais critérios de priorização?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "comercial": {
        "title": "Comercial / Vendas",
        "questions": [
            "Como é feita a previsão de vendas? Qual a acurácia histórica?",
            "Existe processo formal de validação de demanda com a produção?",
            "Como é a política de preços? Existe pricing dinâmico?",
            "Qual o lead time prometido ao cliente vs realizado?",
            "Como é gerenciado o portfólio de clientes (ABC, segmentação)?",
            "Quais são as principais reclamações de clientes sobre entrega e atendimento?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "logistica": {
        "title": "Logística",
        "questions": [
            "Como funciona o fluxo de recebimento, armazenagem e expedição?",
            "Qual o nível de automação na operação (picking, embalagem, inventário)?",
            "Como é gerenciada a frota / transportadoras? Existe TMS?",
            "Qual o custo logístico como % do faturamento?",
            "Existe rastreabilidade de pedidos em tempo real?",
            "Como é feita a roteirização de entregas? Manual ou otimizada?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "ti": {
        "title": "Tecnologia / TI",
        "questions": [
            "Quais sistemas são usados no dia a dia (ERP, WMS, TMS, BI)?",
            "Qual a versão do ERP e quando foi a última atualização?",
            "Existem integrações automáticas entre sistemas ou é tudo manual?",
            "Como é a qualidade e confiabilidade dos dados para tomada de decisão?",
            "Existe algum projeto de BI, data lake ou analytics em andamento?",
            "Como é o suporte de TI para a operação? Qual o SLA?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "financeiro": {
        "title": "Financeiro / Controladoria",
        "questions": [
            "Como é feito o custeio de produtos? Custo padrão ou real?",
            "Qual a visibilidade de margem por SKU, canal e cliente?",
            "Como são aprovados investimentos em supply chain? Qual o processo?",
            "Existe análise de capital de giro vinculada ao estoque?",
            "Como é feito o orçamento anual de operações? É bottom-up ou top-down?",
            "Quais KPIs financeiros são acompanhados com frequência mensal?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "qualidade": {
        "title": "Qualidade",
        "questions": [
            "Como é feito o controle de qualidade na produção? Em quais etapas?",
            "Os checklists de qualidade são digitais ou em papel?",
            "Qual o índice de retrabalho e refugo atual? Como é medido?",
            "Existe sistema de rastreabilidade de lotes/matérias-primas?",
            "Como são tratadas as não-conformidades? Existe processo formal?",
            "A empresa tem certificações ISO? Como é a manutenção?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "compras": {
        "title": "Compras / Procurement",
        "questions": [
            "Como é o processo de compras? Manual, por aprovação ou automatizado?",
            "Existe avaliação formal de fornecedores? Quais critérios?",
            "Qual o lead time médio de compra das principais matérias-primas?",
            "Existe dependência crítica de fornecedor único em algum insumo?",
            "Como é feita a negociação de contratos? Existem contratos de longo prazo?",
            "Qual o saving anual gerado pela área de compras?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "rh": {
        "title": "RH / Pessoas",
        "questions": [
            "Como está estruturada a área de RH? Quantas pessoas compõem a equipe e como estão divididas as responsabilidades?",
            "Existe programa de capacitação técnica para a equipe?",
            "Qual o turnover da área? Quais os cargos mais críticos para reter?",
            "Existe clareza de papéis e responsabilidades (RACI) nos processos-chave?",
            "Como funciona a comunicação entre turnos e entre áreas?",
            "Existem programas de reconhecimento por performance operacional?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
    "diretoria": {
        "title": "Diretoria Geral",
        "questions": [
            "Qual a visão estratégica para supply chain nos próximos 3 anos?",
            "Quais são os principais riscos do negócio relacionados à operação?",
            "Como é priorizado o investimento entre projetos concorrentes?",
            "A empresa tem visão de supply chain digital? Qual o horizonte?",
            "Quais foram as maiores mudanças na operação nos últimos 2 anos?",
            "O que a diretoria considera como 'estado ideal' da operação em 3 anos?",
            "Na sua rotina de trabalho, qual a tarefa que te demanda mais tempo? Qual a maior dificuldade para executá-la?",
        ]
    },
}


def generate_form_docx(area: str, interviewee: str = "", role: str = "", date: str = "") -> bytes:
    """Gera um formulário .docx formatado para uma área específica."""
    area_data = AREA_QUESTIONS.get(area)
    if not area_data:
        raise ValueError(f"Área '{area}' não encontrada")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Header ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STOKEN ADVISORY")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x8B, 0x1D, 0x1D)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Formulário de Entrevista — Diagnóstico Estratégico")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x68, 0x60)

    doc.add_paragraph()  # spacer

    # ── Info table ──
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells = [
        ("ÁREA:", area_data["title"], "DATA:", date or "___/___/______"),
        ("ENTREVISTADO:", interviewee or "________________________________", "CARGO:", role or "________________________________"),
        ("ENTREVISTADOR:", "________________________________", "DURAÇÃO:", "______ min"),
    ]

    for i, (l1, v1, l2, v2) in enumerate(cells):
        row = table.rows[i]
        _set_cell(row.cells[0], l1, bold=True, size=8, color="6B6860")
        _set_cell(row.cells[1], v1, size=10)
        _set_cell(row.cells[2], l2, bold=True, size=8, color="6B6860")
        _set_cell(row.cells[3], v2, size=10)

    # Style table borders
    _style_table(table)

    doc.add_paragraph()  # spacer

    # ── Section title ──
    p = doc.add_paragraph()
    run = p.add_run(f"PERGUNTAS — {area_data['title'].upper()}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x8B, 0x1D, 0x1D)
    run.bold = True

    _add_line(doc)

    # ── Questions ──
    for i, question in enumerate(area_data["questions"], 1):
        # Question
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {question}")
        run.font.size = Pt(11)
        run.bold = True
        p.space_after = Pt(4)

        # Answer area
        p = doc.add_paragraph()
        run = p.add_run("R:")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x9B, 0x98, 0x8F)
        p.space_after = Pt(2)

        # Lines for writing
        for _ in range(4):
            p = doc.add_paragraph()
            run = p.add_run("_" * 85)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE6, 0xE4, 0xDD)
            p.space_after = Pt(0)
            p.space_before = Pt(6)

        doc.add_paragraph()  # spacer between questions

    # ── Additional observations ──
    _add_line(doc)
    p = doc.add_paragraph()
    run = p.add_run("OBSERVAÇÕES ADICIONAIS")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x8B, 0x1D, 0x1D)
    run.bold = True

    for _ in range(8):
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xE6, 0xE4, 0xDD)
        p.space_after = Pt(0)
        p.space_before = Pt(6)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Stoken Advisory — Documento Confidencial")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9B, 0x98, 0x8F)
    run.italic = True

    # Export to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def parse_form_docx(file_bytes: bytes) -> dict:
    """Parseia um formulário .docx preenchido e extrai dados."""
    doc = Document(io.BytesIO(file_bytes))

    result = {
        "interviewee": "",
        "role": "",
        "department": "",
        "date": "",
        "interviewer": "",
        "transcript": "",
        "questions_answered": 0,
        "questions_total": 0,
    }

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    text = "\n".join(full_text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            for i, cell in enumerate(cells):
                if "ÁREA:" in cell and i + 1 < len(cells):
                    result["department"] = cells[i + 1]
                if "ENTREVISTADO:" in cell and i + 1 < len(cells):
                    val = cells[i + 1].replace("_", "").strip()
                    if val:
                        result["interviewee"] = val
                if "CARGO:" in cell and i + 1 < len(cells):
                    val = cells[i + 1].replace("_", "").strip()
                    if val:
                        result["role"] = val
                if "DATA:" in cell and i + 1 < len(cells):
                    val = cells[i + 1].replace("_", "").strip()
                    if val:
                        result["date"] = val
                if "ENTREVISTADOR:" in cell and i + 1 < len(cells):
                    val = cells[i + 1].replace("_", "").strip()
                    if val:
                        result["interviewer"] = val

    # Extract Q&A from paragraphs
    transcript_parts = []
    current_question = None
    current_answer_lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        # Remove underlines
        t_clean = t.replace("_", "").strip()

        # Detect question (numbered: "1. ...", "2. ...")
        q_match = re.match(r'^(\d+)\.\s+(.+)', t)
        if q_match:
            # Save previous Q&A
            if current_question:
                answer = " ".join(current_answer_lines).strip()
                if answer:
                    result["questions_answered"] += 1
                transcript_parts.append(f"P: {current_question}")
                transcript_parts.append(f"R: {answer if answer else '(sem resposta)'}")
                transcript_parts.append("")
                result["questions_total"] += 1

            current_question = q_match.group(2)
            current_answer_lines = []
            continue

        # Detect answer line
        if t.startswith("R:"):
            answer_text = t[2:].strip()
            if answer_text:
                current_answer_lines.append(answer_text)
            continue

        # Continuation of answer (non-empty, not a section header)
        if current_question and t_clean and not t.startswith("PERGUNTAS") and not t.startswith("OBSERV") and not t.startswith("STOKEN") and not t.startswith("Formulário") and not t.startswith("Stoken"):
            current_answer_lines.append(t_clean)

    # Save last Q&A
    if current_question:
        answer = " ".join(current_answer_lines).strip()
        if answer:
            result["questions_answered"] += 1
        transcript_parts.append(f"P: {current_question}")
        transcript_parts.append(f"R: {answer if answer else '(sem resposta)'}")
        result["questions_total"] += 1

    result["transcript"] = "\n".join(transcript_parts)
    return result


def _set_cell(cell, text, bold=False, size=10, color=None):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )


def _style_table(table):
    """Apply minimal borders to table."""
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:color'): 'E6E4DD',
            qn('w:space'): '0',
        })
        borders.append(el)
    tblPr.append(borders)


def _add_line(doc):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    run = p.add_run("─" * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xE6, 0xE4, 0xDD)
    p.space_before = Pt(4)
    p.space_after = Pt(4)
